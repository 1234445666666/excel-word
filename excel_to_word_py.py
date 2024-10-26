import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
# импорт pandas и docx


def set_table_border(table):
    tbl = table._tbl
    tblBorders = parse_xml(r'<w:tblBorders {}>'
                           '<w:top w:val="single" w:sz="4" wSpace="0"/>'
                           '<w:left w:val="single" w:sz="4" wSpace="0"/>'
                           '<w:bottom w:val="single" w:sz="4" wSpace="0"/>'
                           '<w:right w:val="single" w:sz="4" wSpace="0"/>'
                           '<w:insideH w:val="single" w:sz="4" wSpace="0"/>'
                           '<w:insideV w:val="single" w:sz="4" wSpace="0"/>'
                           '</w:tblBorders>'.format(nsdecls('w')))

    # Установка границ
    tbl.tblPr.append(tblBorders)


# путь к excel файлу и его чтение
excel_file_path = 'exlel.xlsx'

xls = pd.ExcelFile(excel_file_path)

# создание нового документа word
doc = Document()

# чтение всех листов в документ
for sheet_name in xls.sheet_names:  # Проходим по всем листам
    df = pd.read_excel(xls, sheet_name=sheet_name)

    # Добавляем заголовок листа
    doc.add_heading(sheet_name, level=1)

    # Создаем таблицу в Word
    table = doc.add_table(rows=1, cols=len(df.columns))
    set_table_border(table)  # Установка границ таблицы

    # Заполняем заголовки таблицы
    for i, column in enumerate(df.columns):
        cell = table.cell(0, i)
        cell.text = str(column)

    # Заполняем данные таблицы
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)

    set_table_border(table)  # Установка границ таблицы


# сохранение документа word
doc.save('output.docx')
